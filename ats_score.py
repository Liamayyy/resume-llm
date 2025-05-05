#!/usr/bin/env python3
"""
ats_score.py  –  Pure‑function ATS résumé scorer
================================================
Export
------
    ats_score(resume_path: str, job_path: str, *, use_tfidf=False) -> int
Returns an integer 0‑100.

All logic from the old ats_scoring.py is inlined here (CLI removed,
printing silenced).  External deps unchanged:
    pip install python-docx pdfplumber sentence-transformers spacy \
                scikit-learn nltk
First run will auto‑download NLTK stop‑words + spaCy “en_core_web_sm”.
"""
from __future__ import annotations
from pathlib import Path
from typing import Dict, List, Tuple
import re, sys

# ─────────────────────── TEXT EXTRACTION ──────────────────────────
def _extract_text(path: Path) -> Tuple[str, Dict[str, bool]]:
    info = {"has_tables": False, "has_images": False}
    suf = path.suffix.lower()

    if suf == ".docx":
        import docx
        doc = docx.Document(path)
        info["has_tables"] = bool(doc.tables)
        text = "\n".join(p.text for p in doc.paragraphs)

    elif suf == ".pdf":
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for p in pdf.pages:
                if p.images:
                    info["has_images"] = True
                if p.extract_tables():
                    info["has_tables"] = True
                pages.append(p.extract_text() or "")
        text = "\n".join(pages)
    else:
        raise ValueError(f"Unsupported file type: {suf}")
    return text, info

# ─────────────────────── NORMALISATION ───────────────────────────
_STOPS: set[str] | None = None
def _normalise(txt: str) -> str:
    global _STOPS
    if _STOPS is None:
        import nltk, nltk.corpus
        try:
            _STOPS = set(nltk.corpus.stopwords.words("english"))
        except LookupError:
            nltk.download("stopwords")
            _STOPS = set(nltk.corpus.stopwords.words("english"))
    txt = txt.lower()
    txt = re.sub(r"[^a-z0-9\s]", " ", txt)
    tokens = [t for t in txt.split() if t not in _STOPS and len(t) > 2]
    return " ".join(tokens)

# ─────────────────────── STRUCTURED PARSE ─────────────────────────
def _structured_parse(raw: str) -> Dict[str, bool]:
    import spacy, re as _re
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        import spacy.cli
        spacy.cli.download("en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")

    doc = nlp(raw)
    return {
        "has_email": bool(_re.search(r"[\w\.-]+@[\w\.-]+\.\w+", raw)),
        "has_phone": bool(_re.search(r"(\+?\d[\d\-\s\(\)]{7,}\d)", raw)),
        "has_name":  any(ent.label_ == "PERSON" for ent in doc.ents),
    }

# ─────────────────────── VECTORS & METRICS ───────────────────────
def _vectorise(texts: List[str], use_tfidf: bool):
    if use_tfidf:
        from sklearn.feature_extraction.text import TfidfVectorizer
        vec = TfidfVectorizer()
        return vec.fit_transform(texts)
    from sentence_transformers import SentenceTransformer
    model = SentenceTransformer("all-MiniLM-L6-v2")
    return model.encode(texts, convert_to_tensor=True, normalize_embeddings=True)

def _cosine(a, b) -> float:
    import numpy as np, torch
    from sklearn.metrics.pairwise import cosine_similarity

    def _np(x):
        if isinstance(x, torch.Tensor):
            return x.detach().cpu().numpy()
        if hasattr(x, "toarray"):          # sparse
            return x.toarray()
        return np.asarray(x)

    return float(cosine_similarity(_np(a), _np(b))[0, 0])

def _keyword_coverage(jd: str, cv: str) -> float:
    jd_set, cv_set = set(jd.split()), set(cv.split())
    return 0.0 if not jd_set else len(jd_set & cv_set) / len(jd_set)

# ─────────────────────── PUBLIC API ──────────────────────────────
def ats_score(resume_path: str, job_path: str, *, use_tfidf: bool = False) -> int:
    jd_raw, _ = _extract_text(Path(job_path))
    cv_raw, fmt = _extract_text(Path(resume_path))

    flags = _structured_parse(cv_raw)          # contact info etc.
    jd_clean = _normalise(jd_raw)
    cv_clean = _normalise(cv_raw)

    vecs = _vectorise([jd_clean, cv_clean], use_tfidf)
    sim  = _cosine(vecs[0:1], vecs[1:2])       # 0‑1
    cov  = _keyword_coverage(jd_clean, cv_clean)

    base = 0.70 * sim + 0.30 * cov
    score = int(round(base * 100))

    compliance_ok = all(flags.values()) and not (fmt["has_tables"] or fmt["has_images"])
    if not compliance_ok:
        score //= 2
    return score

# quick CLI for ad‑hoc tests --------------------------------------
if __name__ == "__main__":
    import argparse
    p = argparse.ArgumentParser(description="Return ATS score 0‑100")
    p.add_argument("-j", "--jd", required=True)
    p.add_argument("-r", "--resume", required=True)
    p.add_argument("--tfidf", action="store_true")
    args = p.parse_args()
    print(ats_score(args.resume, args.jd, use_tfidf=args.tfidf))